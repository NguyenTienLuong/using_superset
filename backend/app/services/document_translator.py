import base64
import io
import os
import tempfile
from typing import Callable, Tuple, Coroutine

from docx import Document
from pdf2docx import Converter


def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """
    Converts a PDF file (in bytes) to a DOCX file (in bytes) using pdf2docx.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as pdf_file:
        pdf_file.write(pdf_bytes)
        pdf_path = pdf_file.name
        
    docx_path = pdf_path + ".docx"
    
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        with open(docx_path, "rb") as docx_file:
            docx_bytes = docx_file.read()
            
        return docx_bytes
    finally:
        if os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
            except Exception:
                pass
        if os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except Exception:
                pass


async def translate_docx_document(
    file_bytes: bytes,
    translate_batch_fn: Callable[[list[str]], Coroutine[None, None, list[str]]]
) -> Tuple[str, str]:
    """
    Translates a DOCX document run by run, preserving the layout using batch translation.
    Returns:
        Tuple[str, str]: (all_translated_text, base64_encoded_docx)
    """
    doc = Document(io.BytesIO(file_bytes))
    
    runs_to_process = []
    
    def collect_paragraph(paragraph):
        if paragraph.text.strip():
            for run in paragraph.runs:
                if run.text.strip():
                    leading_spaces = run.text[:len(run.text) - len(run.text.lstrip())]
                    trailing_spaces = run.text[len(run.text.rstrip()):]
                    original_clean = run.text.strip()
                    runs_to_process.append((paragraph, run, leading_spaces, trailing_spaces, original_clean))

    for paragraph in doc.paragraphs:
        collect_paragraph(paragraph)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    collect_paragraph(paragraph)
                    
    if runs_to_process:
        texts = [r[4] for r in runs_to_process]
        # Await the async translation function
        translated_texts = await translate_batch_fn(texts)
        
        for i, (paragraph, run, leading_spaces, trailing_spaces, original_clean) in enumerate(runs_to_process):
            translated_clean = translated_texts[i] if i < len(translated_texts) else original_clean
            if not translated_clean:
                translated_clean = original_clean
            
            if original_clean.isupper():
                translated_clean = translated_clean.upper()
            elif original_clean.istitle():
                translated_clean = translated_clean.title()
            elif original_clean and original_clean[0].isupper():
                translated_clean = translated_clean[0].upper() + translated_clean[1:]
            elif original_clean.islower():
                translated_clean = translated_clean.lower()
                
            run.text = leading_spaces + translated_clean + trailing_spaces

    all_translated_texts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            all_translated_texts.append(paragraph.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        all_translated_texts.append(paragraph.text.strip())
                        
    # Save to BytesIO
    out_stream = io.BytesIO()
    doc.save(out_stream)
    
    # Encode to base64
    b64_data = base64.b64encode(out_stream.getvalue()).decode("utf-8")
    
    full_text = "\n\n".join(all_translated_texts)
    return full_text, b64_data


async def translate_txt_document(
    file_bytes: bytes,
    translate_fn: Callable[[str], Coroutine[None, None, str]]
) -> Tuple[str, str]:
    """
    Translates a TXT document line by line, preserving newlines.
    Returns:
        Tuple[str, str]: (translated_text, base64_encoded_txt)
    """
    original_text = file_bytes.decode("utf-8")
    lines = original_text.split("\n")
    
    translated_lines = []
    for line in lines:
        stripped_line = line.strip()
        if stripped_line:
            # Await the async translation function
            translated_lines.append(await translate_fn(stripped_line))
        else:
            translated_lines.append("")  # Preserve empty lines
            
    translated_text = "\n".join(translated_lines)
    b64_data = base64.b64encode(translated_text.encode("utf-8")).decode("utf-8")
    
    return translated_text, b64_data
